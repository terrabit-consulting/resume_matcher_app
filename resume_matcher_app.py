# app.py
import openai
import streamlit as st
import time
import fitz  # PyMuPDF
import docx
import pandas as pd
import re
import io
import tempfile
import os

# -------- Optional .doc support (best-effort) ----------
try:
    import textract  # requires antiword/catdoc on system for .doc
    HAS_TEXTRACT = True
except Exception:
    HAS_TEXTRACT = False

# ---------------- OpenAI client ------------------------
client = openai.OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

# ---------------- Constants/Regex ----------------------
EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
SCORE_RE = re.compile(r"(?i)\bscore\b[^0-9]{0,12}(\d{1,3})\s*%")

# ---------------- Utilities ---------------------------
def trim(txt: str, max_chars: int = 12000) -> str:
    """Hard cap to control tokens."""
    if not txt:
        return ""
    return txt[:max_chars]

def call_gpt_with_fallback(prompt: str) -> str:
    """Retry with backoff + fallback model for reliability."""
    models = ["gpt-4o", "gpt-4o-mini", "gpt-4o-mini"]
    for attempt, model in enumerate(models):
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You are a concise recruiter assistant."},
                    {"role": "user", "content": prompt},
                ],
                temperature=0,
            )
            return resp.choices[0].message.content.strip()
        except Exception as e:
            if attempt == len(models) - 1:
                st.error(f"❌ OpenAI error: {e}")
            time.sleep(1.5 * (attempt + 1))
    return "⚠️ GPT processing failed."

# ---------------- File readers ------------------------
def read_pdf(file) -> str:
    text = ""
    with fitz.open(stream=file.read(), filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text

def read_docx(file) -> str:
    d = docx.Document(file)
    parts = []
    for p in d.paragraphs:
        parts.append(p.text)
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    try:
        section = d.sections[0]
        for p in section.footer.paragraphs:
            parts.append(p.text)
    except Exception:
        pass
    return "\n".join(parts)

def read_doc(file, name: str) -> str:
    """Best-effort .doc using textract; otherwise advice."""
    if not HAS_TEXTRACT:
        st.warning(f"⚠️ {name}: .doc reading needs `textract` (+ antiword/catdoc). "
                   "Please convert to .docx/.pdf or install deps.")
        return ""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp:
        tmp.write(file.read())
        tmp_path = tmp.name
    try:
        raw = textract.process(tmp_path)
        return raw.decode("utf-8", errors="ignore")
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass

def read_file(file):
    # Some browsers send generic MIME types; also check extension
    name = getattr(file, "name", "").lower()
    mime = getattr(file, "type", "")
    if mime == "application/pdf" or name.endswith(".pdf"):
        return read_pdf(file)
    if (mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document") or name.endswith(".docx"):
        return read_docx(file)
    if mime == "application/msword" or name.endswith(".doc"):
        return read_doc(file, name)
    # fallback plain text
    return file.read().decode("utf-8", errors="ignore")

# ---------------- Parsers -----------------------------
def extract_email(text: str) -> str:
    m = EMAIL_RE.search(text or "")
    return m.group(0) if m else "Not found"

def extract_candidate_name_from_table(text: str):
    matches = re.findall(r"(?i)Candidate Name\s*[\t:–-]*\s*(.+)", text or "")
    for match in matches:
        name = match.strip()
        if 2 <= len(name.split()) <= 4:
            return name.title()
    return None

def extract_candidate_name_from_footer(text: str):
    m = re.search(r"(?i)Resume of\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})", text or "")
    if m:
        return m.group(1).strip().title()
    return None

def improved_extract_candidate_name(text: str, filename: str):
    """GPT last resort on a trimmed header slice."""
    trimmed_text = "\n".join((text or "").splitlines()[:50])
    prompt = f"""
Extract the candidate's full name ONLY from the resume text below.
Return only the name; if unsure, return: Name Not Found

Resume:
{trimmed_text}
"""
    name = call_gpt_with_fallback(prompt)
    suspicious = ["java","python","developer","resume","engineer","servers","experience","summary"]
    if (
        not name
        or len(name.split()) > 5
        or any(w in name.lower() for w in suspicious)
        or "@" in name
        or name.lower().startswith("name not found")
    ):
        return "Name Not Found"
    return name.strip().title()

def extract_candidate_name(text: str, filename: str):
    return (
        extract_candidate_name_from_table(text)
        or extract_candidate_name_from_footer(text)
        or improved_extract_candidate_name(text, filename)
    )

# ---------------- GPT prompts -------------------------
def compare_resume(jd_text: str, resume_text: str, candidate_name: str) -> str:
    """Forces the first-screenshot block format."""
    prompt = f"""
Compare this resume to the JD and reply ONLY in this exact format:

📛 {candidate_name}
✅ Score: <0–100>%
🔍 Reason:
- 1–5 short bullets on role fit, core skills/tools (Selenium, SQL/Oracle, Linux, CI/CD), domains, deployments/logs

<If score < 70 include this exact line>
⚠️ This candidate may not meet the requirements.

JD:
{trim(jd_text)}

Resume:
{trim(resume_text)}
"""
    return call_gpt_with_fallback(prompt)

def generate_followup(jd_text: str, resume_text: str):
    prompt = f"""
Create 3 things based on this JD + resume:

1) WhatsApp (casual, 3–4 lines)
2) Email (formal, concise, with role, next steps, availability)
3) 3 screening questions focused on gaps vs JD

JD:
{trim(jd_text)}

Resume:
{trim(resume_text)}
"""
    return call_gpt_with_fallback(prompt)

# ---------------- Streamlit UI ------------------------
st.set_page_config(page_title="Resume Matcher GPT", layout="centered")
st.title("📌 Terrabit Consulting Talent Match System")
st.write("Upload a JD and multiple resumes. Get match scores, red flags, and follow-up messaging.")

# Session state
for key, default in [
    ("results", []),
    ("processed_resumes", set()),
    ("jd_text", ""),
    ("jd_file", None),
    ("summary", []),
]:
    if key not in st.session_state:
        st.session_state[key] = default

if st.button("🔁 Start New Matching Session"):
    st.session_state.clear()
    st.rerun()

jd_file = st.file_uploader("📄 Upload Job Description", type=["txt", "pdf", "docx", "doc"], key="jd_uploader")
resume_files = st.file_uploader("📑 Upload Candidate Resumes", type=["txt", "pdf", "docx", "doc"], accept_multiple_files=True, key="resume_uploader")

# Load JD once
if jd_file and not st.session_state.get("jd_text"):
    jd_text = read_file(jd_file)
    st.session_state["jd_text"] = jd_text
    st.session_state["jd_file"] = getattr(jd_file, "name", "JD")

jd_text = st.session_state.get("jd_text", "")

# Guard UI
run_disabled = not (jd_text and resume_files)
st.button("🚀 Run Matching", disabled=run_disabled, key="run_guard")

if not run_disabled and st.session_state.get("run_clicked") is None:
    st.session_state["run_clicked"] = True
    for resume_file in resume_files:
        if resume_file.name in st.session_state["processed_resumes"]:
            continue

        resume_text = read_file(resume_file)
        if not resume_text:
            # Already warned in read_file for .doc without textract
            continue

        correct_name = extract_candidate_name(resume_text, resume_file.name)
        correct_email = extract_email(resume_text)

        with st.spinner(f"🔎 Analyzing {correct_name}..."):
            result = compare_resume(jd_text, resume_text, correct_name)

        m = SCORE_RE.search(result or "")
        score = max(0, min(100, int(m.group(1)))) if m else 0

        st.session_state["results"].append({
            "correct_name": correct_name,
            "email": correct_email,
            "score": score,
            "result": result,
            "resume_text": resume_text
        })
        st.session_state["processed_resumes"].add(resume_file.name)

        st.session_state["summary"].append({
            "Candidate Name": correct_name,
            "Email": correct_email,
            "Score": score
        })

# Render results in first‑screenshot style
for entry in st.session_state["results"]:
    st.markdown("---")
    st.subheader(f"📛 {entry['correct_name']}")
    st.markdown(f"📧 **Email**: {entry['email']}")
    st.markdown(entry["result"], unsafe_allow_html=True)

    score = entry["score"]
    if score < 50:
        st.error("❌ Not suitable – Major role mismatch")
    elif score < 70:
        st.warning("⚠️ Consider with caution – Lacks core skills")
    else:
        st.success("✅ Strong match – Good alignment with JD")

    if st.button(f"✉️ Generate Follow-up for {entry['correct_name']}", key=f"followup_{entry['correct_name']}"):
        with st.spinner("Generating messages..."):
            followup = generate_followup(jd_text, entry["resume_text"])
            st.markdown("---")
            st.markdown(followup, unsafe_allow_html=True)

# Summary + download
if st.session_state["summary"]:
    st.markdown("### 📊 Summary of All Candidates")
    df_summary = pd.DataFrame(st.session_state["summary"]).sort_values(by="Score", ascending=False)
    st.dataframe(df_summary, use_container_width=True)

    excel_buffer = io.BytesIO()
    with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
        df_summary.to_excel(writer, index=False)

    st.download_button(
        label="📥 Download Summary as Excel",
        data=excel_buffer.getvalue(),
        file_name="resume_match_summary.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
